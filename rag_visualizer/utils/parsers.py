"""Document parsing utilities for various file formats.

Supports parsing PDF, TXT, MD, and DOCX files into plain text.
"""

import os
from dataclasses import dataclass
from pathlib import Path

import streamlit as st
from PIL import Image as PILImage

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import markdown
except ImportError:
    markdown = None

# These are now required dependencies
from docling.backend.docling_parse_v4_backend import DoclingParseV4DocumentBackend
from docling.datamodel.accelerator_options import AcceleratorOptions, AcceleratorDevice
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import TableStructureOptions
from docling.document_converter import DocumentConverter, FormatOption
from docling.pipeline.standard_pdf_pipeline import (
    StandardPdfPipeline,
    ThreadedPdfPipelineOptions,
)
from docling_core.types.doc import DocItemLabel, PictureItem

# Mapping of file extensions to Docling InputFormat
DOCLING_FORMAT_MAP = {
    ".pdf": InputFormat.PDF,
    ".docx": InputFormat.DOCX,
    ".pptx": InputFormat.PPTX,
    ".xlsx": InputFormat.XLSX,
    ".html": InputFormat.HTML,
    ".htm": InputFormat.HTML,
    ".png": InputFormat.IMAGE,
    ".jpg": InputFormat.IMAGE,
    ".jpeg": InputFormat.IMAGE,
    ".bmp": InputFormat.IMAGE,
    ".tiff": InputFormat.IMAGE,
    ".tif": InputFormat.IMAGE,
}


@dataclass
class ExtractedImage:
    """An image extracted from a document."""
    index: int  # 1-based index
    pil_image: PILImage.Image
    caption: str | None = None
    page_number: int | None = None


@st.cache_resource(show_spinner=False)
def _get_docling_converter(
    enable_ocr: bool,
    enable_table_structure: bool,
    num_threads: int,
    device: str = "auto",
) -> DocumentConverter:
    """Create a cached docling converter tuned for speed."""
    safe_threads = max(1, int(num_threads or 4))

    pipeline_options = ThreadedPdfPipelineOptions(
        accelerator_options=AcceleratorOptions(
            num_threads=safe_threads,
            device=AcceleratorDevice.AUTO,
        ),
        do_ocr=enable_ocr,
        do_table_structure=enable_table_structure,
        table_structure_options=TableStructureOptions(enable_table_matching=True, enable_table_merging=True, enable_table_splitting=True, enable_table_reconstruction=True),
        # Disable non-essential outputs for speed
        generate_page_images=False,
        generate_table_images=False,
        generate_picture_images=False,
        generate_parsed_pages=False,
        do_picture_classification=False,
        do_picture_description=False,
    )

    format_option = FormatOption(
        pipeline_options=pipeline_options,
        backend=DoclingParseV4DocumentBackend,
        pipeline_cls=StandardPdfPipeline,
    )

    return DocumentConverter(
        allowed_formats=[InputFormat.PDF],
        format_options={InputFormat.PDF: format_option},
    )


@st.cache_resource(show_spinner=False)
def _get_generic_docling_converter(input_format: InputFormat) -> DocumentConverter:
    """Create a cached docling converter for non-PDF formats.

    Args:
        input_format: The InputFormat to convert (DOCX, PPTX, XLSX, HTML, IMAGE)

    Returns:
        Configured DocumentConverter for the specified format
    """
    return DocumentConverter(allowed_formats=[input_format])


@st.cache_resource(show_spinner=False)
def _get_docling_converter_with_images(
    enable_ocr: bool,
    enable_table_structure: bool,
    num_threads: int,
    device: str = "auto",
) -> DocumentConverter:
    """Create a cached docling converter with image extraction enabled."""
    safe_threads = max(1, int(num_threads or 4))

    pipeline_options = ThreadedPdfPipelineOptions(
        accelerator_options=AcceleratorOptions(
            num_threads=safe_threads,
            device=AcceleratorDevice.AUTO,
        ),
        do_ocr=enable_ocr,
        do_table_structure=enable_table_structure,
        table_structure_options=TableStructureOptions(
            enable_table_matching=True,
            enable_table_merging=True,
            enable_table_splitting=True,
            enable_table_reconstruction=True,
        ),
        # Enable image extraction
        generate_page_images=True,
        generate_table_images=False,
        generate_picture_images=True,
        generate_parsed_pages=False,
        do_picture_classification=False,
        do_picture_description=False,
    )

    format_option = FormatOption(
        pipeline_options=pipeline_options,
        backend=DoclingParseV4DocumentBackend,
        pipeline_cls=StandardPdfPipeline,
    )

    return DocumentConverter(
        allowed_formats=[InputFormat.PDF],
        format_options={InputFormat.PDF: format_option},
    )


def _extract_images_from_docling(doc: object) -> list[ExtractedImage]:
    """Extract images from a docling document result."""
    images = []
    index = 1

    for item, _level in doc.iterate_items():
        if isinstance(item, PictureItem):
            try:
                pil_image = item.get_image(doc)
                if pil_image is not None:
                    # Try to get page number from item's provenance
                    page_number = None
                    if hasattr(item, 'prov') and item.prov:
                        for prov in item.prov:
                            if hasattr(prov, 'page_no'):
                                page_number = prov.page_no
                                break

                    images.append(ExtractedImage(
                        index=index,
                        pil_image=pil_image,
                        caption=None,
                        page_number=page_number,
                    ))
                    index += 1
            except Exception:
                # Skip images that fail to extract
                continue

    return images


def _filter_docling_items(doc: object, filter_labels: list[DocItemLabel]) -> str:
    """Export docling document to markdown with specific items filtered out.
    
    Args:
        doc: Docling document object
        filter_labels: List of DocItemLabel types to filter out
        
    Returns:
        Markdown text with filtered items removed
    """
    if not filter_labels:
        # No filtering needed, use default export
        return doc.export_to_markdown()
    
    # Manually build markdown by iterating through items and skipping filtered ones
    
    markdown_parts = []
    
    for item, _level in doc.iterate_items():
        # Skip items with filtered labels
        if hasattr(item, 'label') and item.label in filter_labels:
            continue
        
        # Convert item to markdown text
        try:
            # Get the text representation of the item
            if hasattr(item, 'export_to_markdown'):
                item_markdown = item.export_to_markdown()
                if item_markdown and item_markdown.strip():
                    markdown_parts.append(item_markdown)
            elif hasattr(item, 'text'):
                # For simple text items
                if item.text and item.text.strip():
                    markdown_parts.append(item.text)
        except Exception:
            # Skip items that fail to export
            continue
    
    return "\n\n".join(markdown_parts) if markdown_parts else doc.export_to_markdown()


def export_document(doc: object, output_format: str, filter_labels: list | None = None) -> str:
    """Export a Docling document to the specified format.

    Args:
        doc: Docling document object
        output_format: Target format (markdown, html, doctags, json)
        filter_labels: Optional list of DocItemLabel types to filter out (markdown only)

    Returns:
        Exported document text in the specified format
    """
    if output_format == "markdown":
        # Use filtering for markdown export
        if filter_labels:
            return _filter_docling_items(doc, filter_labels)
        return doc.export_to_markdown()
    elif output_format == "html":
        return doc.export_to_html()
    elif output_format == "doctags":
        return doc.export_to_document_tokens()
    elif output_format == "json":
        return doc.model_dump_json()
    else:
        # Default to markdown
        if filter_labels:
            return _filter_docling_items(doc, filter_labels)
        return doc.export_to_markdown()


def parse_pdf_docling(
    content: bytes, params: dict | None = None
) -> tuple[str, list[ExtractedImage]]:
    """Parse PDF content using Docling engine.

    Args:
        content: PDF file content as bytes
        params: Optional parsing configuration including:
            - output_format: Export format (markdown, html, doctags, json)
            - docling_enable_ocr: Enable OCR for scanned PDFs
            - docling_table_structure: Extract table structure
            - docling_threads: Number of worker threads
            - docling_device: Compute device (auto, cpu, cuda, mps)
            - docling_filter_labels: Labels to filter out (markdown only)
            - docling_extract_images: Extract images from PDF

    Returns:
        Tuple of (exported_text, extracted_images)

    Raises:
        ValueError: If PDF parsing fails
    """
    params = params or {}

    # For large digital PDFs we skip OCR by default for speed, but allow opt-in.
    enable_ocr = bool(params.get("docling_enable_ocr", False))
    enable_table_structure = bool(params.get("docling_table_structure", True))
    num_threads = params.get("docling_threads") or os.cpu_count() or 4
    extract_images = bool(params.get("docling_extract_images", False))
    device = params.get("docling_device", "auto")
    
    # Get output format (markdown, html, doctags, json)
    output_format = params.get("output_format", "markdown")

    # Build list of labels to filter from the list of label names (markdown only)
    filter_label_names = params.get("docling_filter_labels", [])
    filter_labels = []
    for label_name in filter_label_names:
        try:
            filter_labels.append(DocItemLabel(label_name.lower()))
        except ValueError:
            # Skip invalid label names
            pass

    try:
        import tempfile

        # docling requires a file path, so we'll use a temp file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            # Use image-enabled converter if image extraction is requested
            if extract_images:
                converter = _get_docling_converter_with_images(
                    enable_ocr=enable_ocr,
                    enable_table_structure=enable_table_structure,
                    num_threads=num_threads,
                    device=device,
                )
            else:
                converter = _get_docling_converter(
                    enable_ocr=enable_ocr,
                    enable_table_structure=enable_table_structure,
                    num_threads=num_threads,
                    device=device,
                )

            result = converter.convert(tmp_path)
            doc = result.document

            # Export document in the specified format
            text = export_document(doc, output_format, filter_labels)

            # Extract images if enabled
            images: list[ExtractedImage] = []
            if extract_images:
                images = _extract_images_from_docling(doc)

            return text, images
        finally:
            # Clean up temp file
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass
    except Exception as e:
        raise ValueError(f"Failed to parse PDF with docling: {str(e)}") from e


def parse_pdf(
    content: bytes, params: dict | None = None
) -> tuple[str, list[ExtractedImage]]:
    """Parse PDF content using Docling.

    Args:
        content: PDF file content as bytes
        params: Optional parsing configuration

    Returns:
        Tuple of (text_content, extracted_images)

    Raises:
        ValueError: If PDF parsing fails
    """
    return parse_pdf_docling(content, params)


def parse_with_docling(
    content: bytes,
    extension: str,
    params: dict | None = None,
) -> tuple[str, list[ExtractedImage]]:
    """Parse document content using Docling for supported formats.

    Supports PPTX, XLSX, HTML, and image formats.

    Args:
        content: File content as bytes
        extension: File extension (e.g., ".pptx", ".xlsx")
        params: Optional parsing configuration including:
            - output_format: Export format (markdown, html, doctags, json)
            - docling_filter_labels: Labels to filter out (markdown only)
            - docling_extract_images: Extract images from document

    Returns:
        Tuple of (exported_text, extracted_images)

    Raises:
        ValueError: If format is not supported or parsing fails
    """
    params = params or {}
    extension = extension.lower()

    if extension not in DOCLING_FORMAT_MAP:
        raise ValueError(f"Unsupported format for Docling: {extension}")

    input_format = DOCLING_FORMAT_MAP[extension]

    # Get output format (markdown, html, doctags, json)
    output_format = params.get("output_format", "markdown")

    # Build list of labels to filter from the list of label names (markdown only)
    filter_label_names = params.get("docling_filter_labels", [])
    filter_labels = []
    for label_name in filter_label_names:
        try:
            filter_labels.append(DocItemLabel(label_name.lower()))
        except ValueError:
            pass

    try:
        import tempfile

        # Docling requires a file path
        with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            converter = _get_generic_docling_converter(input_format)
            result = converter.convert(tmp_path)
            doc = result.document

            # Export document in the specified format
            text = export_document(doc, output_format, filter_labels)

            # Extract images if present
            images: list[ExtractedImage] = []
            if params.get("docling_extract_images", False):
                images = _extract_images_from_docling(doc)

            return text, images
        finally:
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass
    except Exception as e:
        format_name = extension.lstrip(".").upper()
        raise ValueError(f"Failed to parse {format_name} with Docling: {str(e)}") from e


def _extract_table_as_markdown(table: object) -> str:
    """Extract DOCX table as markdown-style text.

    Args:
        table: python-docx Table object

    Returns:
        Markdown-formatted table text
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))

    # Add separator after header row
    if rows:
        header = rows[0]
        separator = " | ".join(["---"] * len(rows[0].split(" | ")))
        return f"{header}\n{separator}\n" + "\n".join(rows[1:])
    return ""


def parse_docx(content: bytes, params: dict | None = None) -> str:
    """Parse DOCX content with optional markdown conversion.

    Args:
        content: DOCX file content as bytes
        params: Optional parsing configuration

    Returns:
        Extracted text content (optionally as markdown)

    Raises:
        ImportError: If python-docx is not installed
        ValueError: If DOCX parsing fails
    """
    params = params or {}
    if Document is None:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install with: pip install python-docx"
        )

    try:
        from io import BytesIO

        docx_file = BytesIO(content)
        doc = Document(docx_file)
        text_parts = []

        # Process paragraphs
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            # Detect heading styles and convert to markdown
            style_name = paragraph.style.name if paragraph.style else ""

            if "Heading 1" in style_name:
                text_parts.append(f"\n# {paragraph.text}\n")
            elif "Heading 2" in style_name:
                text_parts.append(f"\n## {paragraph.text}\n")
            elif "Heading 3" in style_name:
                text_parts.append(f"\n### {paragraph.text}\n")
            elif "Heading 4" in style_name:
                text_parts.append(f"\n#### {paragraph.text}\n")
            elif "Heading 5" in style_name:
                text_parts.append(f"\n##### {paragraph.text}\n")
            elif "Heading 6" in style_name:
                text_parts.append(f"\n###### {paragraph.text}\n")
            else:
                # Regular paragraph
                text_parts.append(paragraph.text)

        # Process tables if enabled
        if params.get("extract_tables", True) and doc.tables:
            for table in doc.tables:
                table_text = _extract_table_as_markdown(table)
                if table_text:
                    text_parts.append(f"\n\n{table_text}\n")

        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}") from e


def parse_markdown(content: bytes, params: dict | None = None) -> str:
    """Parse Markdown content.

    Args:
        content: Markdown file content as bytes
        params: Optional parsing configuration

    Returns:
        Extracted text content (raw markdown or plain text based on params)

    Raises:
        ImportError: If markdown is not installed
        ValueError: If Markdown parsing fails
    """
    params = params or {}

    try:
        text = content.decode("utf-8")

        # If output format is markdown or original, return raw markdown
        output_format = params.get("output_format", "original")
        if output_format in ["markdown", "original"]:
            return text

        # Otherwise convert to plain text
        if markdown is None:
            # If markdown library not available, just return text as-is
            return text

        # Convert markdown to HTML, then strip HTML tags for plain text
        html = markdown.markdown(text)
        # Simple HTML tag removal (basic implementation)
        import re

        plain_text = re.sub(r"<[^>]+>", "", html)
        # Decode HTML entities
        import html as html_lib

        return html_lib.unescape(plain_text)
    except Exception as e:
        raise ValueError(f"Failed to parse Markdown: {str(e)}") from e


def parse_text(content: bytes) -> str:
    """Parse plain text content.

    Args:
        content: Text file content as bytes

    Returns:
        Text content as string
    """
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        # Try other common encodings
        for encoding in ["latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Failed to decode text file with common encodings") from None


def _normalize_whitespace(text: str) -> str:
    """Normalize excessive whitespace in text.

    Args:
        text: Input text

    Returns:
        Text with normalized whitespace
    """
    import re

    # Replace multiple spaces with single space
    text = re.sub(r" +", " ", text)
    # Replace multiple newlines with double newline
    text = re.sub(r"\n\n+", "\n\n", text)
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(lines)


def _remove_special_chars(text: str) -> str:
    """Remove special characters while preserving basic punctuation.

    Args:
        text: Input text

    Returns:
        Text with special characters removed
    """
    import re

    # Keep letters, numbers, spaces, and basic punctuation (.,-!?:; and -)
    # Hyphen must be escaped/placed last to avoid regex range errors.
    text = re.sub(r"[^\w\s.,!?:;\n-]", "", text)
    return text


def _markdown_to_plain_text(text: str) -> str:
    """Convert markdown content to plain text."""
    import html as html_lib
    import re

    # If markdown lib available, strip tags via HTML conversion for better fidelity
    if markdown is not None:
        html_text = markdown.markdown(text)
        plain_text = re.sub(r"<[^>]+>", "", html_text)
        plain_text = re.sub(r"\s+\n", "\n", plain_text)
        return html_lib.unescape(plain_text)

    # Fallback: lightweight regex cleanup
    text = re.sub(r"^#+\s", "", text, flags=re.MULTILINE)  # Remove headings
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # Remove bold
    text = re.sub(r"_([^_]+)_", r"\1", text)  # Remove italic
    text = re.sub(r"\|[^\n]+\|", "", text)  # Remove table rows
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)  # Remove separators
    text = re.sub(r"^\s*[-*]\s+", "", text, flags=re.MULTILINE)  # Remove list bullets
    return html_lib.unescape(text)


def _convert_to_format(
    text: str, source_format: str, target_format: str  # noqa: ARG001
) -> str:
    """Convert text between different formats.

    For Docling-parsed documents (PDF, PPTX, XLSX, HTML, Image), the output
    format is already applied during parsing via export_document().

    For non-Docling parsers (DOCX, Markdown, Text), this function handles
    format conversion where possible.

    Args:
        text: Input text
        source_format: Source format (PDF, DOCX, Markdown, Text, etc.)
        target_format: Target format (markdown, html, doctags, json)

    Returns:
        Converted text
    """
    # New Docling formats - text is already in the correct format from parsing
    if target_format in ("markdown", "html", "doctags", "json"):
        return text
    # Legacy formats for backwards compatibility
    elif target_format == "original":
        return text
    elif target_format == "plain_text":
        return _markdown_to_plain_text(text)
    else:
        return text


def insert_image_placeholders(text: str, images: list[ExtractedImage]) -> str:
    """Insert image placeholders into document text for RAG indexing.

    For images with captions: [Image N: caption]
    For images without captions: [Image N]

    Args:
        text: Document text
        images: List of extracted images

    Returns:
        Text with image placeholders appended
    """
    if not images:
        return text

    placeholders = []
    for img in images:
        if img.caption:
            placeholders.append(f"[Image {img.index}: {img.caption}]")
        else:
            placeholders.append(f"[Image {img.index}]")

    # Append image placeholders at the end in a dedicated section
    image_section = "\n\n---\n\n## Document Images\n\n" + "\n\n".join(placeholders)
    return text + image_section


def parse_document(
    filename: str, content: bytes, params: dict | None = None
) -> tuple[str, str, list[ExtractedImage]]:
    """Parse a document based on its file extension.

    Args:
        filename: Name of the file (used to determine format)
        content: File content as bytes
        params: Optional parsing configuration dictionary

    Returns:
        Tuple of (parsed_text, file_format, extracted_images)

    Raises:
        ValueError: If file format is not supported or parsing fails
    """
    params = params or {}
    file_path = Path(filename)
    extension = file_path.suffix.lower()

    # Parse document based on format
    images: list[ExtractedImage] = []

    if extension == ".pdf":
        parsed_text, images = parse_pdf(content, params)
        file_format = "PDF"
    elif extension == ".docx":
        parsed_text = parse_docx(content, params)
        file_format = "DOCX"
    elif extension == ".pptx":
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "PPTX"
    elif extension == ".xlsx":
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "XLSX"
    elif extension in [".html", ".htm"]:
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "HTML"
    elif extension in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]:
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "Image"
    elif extension in [".md", ".markdown"]:
        parsed_text = parse_markdown(content, params)
        file_format = "Markdown"
    elif extension == ".txt":
        parsed_text = parse_text(content)
        file_format = "Text"
    else:
        raise ValueError(
            f"Unsupported file format: {extension}. "
            f"Supported formats: .pdf, .docx, .pptx, .xlsx, .html, .md, .txt, .png, .jpg"
        )

    # Apply output format conversion (for non-Docling parsers)
    output_format = params.get("output_format", "markdown")
    parsed_text = _convert_to_format(parsed_text, file_format, output_format)

    # Apply post-processing
    if params.get("normalize_whitespace", False):
        parsed_text = _normalize_whitespace(parsed_text)

    if params.get("remove_special_chars", False):
        parsed_text = _remove_special_chars(parsed_text)

    # Insert image placeholders into text if we have images
    if images:
        parsed_text = insert_image_placeholders(parsed_text, images)

    # Apply global character cap if provided
    max_chars = params.get("max_characters")
    if isinstance(max_chars, int) and max_chars > 0 and len(parsed_text) > max_chars:
        parsed_text = parsed_text[:max_chars]

    return parsed_text, file_format, images


def get_file_preview(text: str, max_length: int = 500) -> str:
    """Get a preview of document text.

    Args:
        text: Full document text
        max_length: Maximum length of preview in characters

    Returns:
        Preview text (truncated if necessary)
    """
    if len(text) <= max_length:
        return text

    # Try to truncate at a sentence boundary
    preview = text[:max_length]
    last_period = preview.rfind(".")
    last_newline = preview.rfind("\n")

    # Use the later of period or newline, or just truncate
    truncate_at = max(last_period, last_newline)
    if truncate_at > max_length * 0.7:  # Only use boundary if it's not too early
        preview = preview[: truncate_at + 1]
    else:
        preview = preview[:max_length] + "..."

    return preview + "..." if len(text) > max_length else preview

