"""Document parsing utilities for various file formats.

Supports parsing PDF, TXT, MD, and DOCX files into plain text.
"""

import inspect
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast

import streamlit as st
from PIL import Image as PILImage

Document: Any | None
try:
    from docx import Document as DocxDocument

    Document = DocxDocument
except ImportError:
    Document = None

markdown: Any | None
try:
    import markdown  # type: ignore[import-untyped]
except ImportError:
    markdown = None

# These are now required dependencies
from docling.backend.docling_parse_v4_backend import DoclingParseV4DocumentBackend
from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import TableStructureOptions
from docling.document_converter import DocumentConverter, FormatOption

try:
    from docling.document_converter import PageRange
except Exception:
    PageRange = None
from docling.pipeline.standard_pdf_pipeline import (  # type: ignore[attr-defined]
    StandardPdfPipeline,
    ThreadedPdfPipelineOptions,
)
from docling_core.types.doc import (  # type: ignore[attr-defined]
    DocItemLabel,
    PictureItem,
)

try:
    from docling_core.transforms.serializer.html import HTMLDocSerializer
    from docling_core.transforms.serializer.markdown import MarkdownDocSerializer
except Exception:
    HTMLDocSerializer = None
    MarkdownDocSerializer = None

# Mapping of file extensions to Docling InputFormat
DOCLING_FORMAT_MAP = {
    ".pdf": InputFormat.PDF,
    ".docx": InputFormat.DOCX,
    ".pptx": InputFormat.PPTX,
    ".xlsx": InputFormat.XLSX,
    ".html": InputFormat.HTML,
    ".htm": InputFormat.HTML,
    ".png": InputFormat.IMAGE,
    ".jpg": InputFormat.IMAGE,
    ".jpeg": InputFormat.IMAGE,
    ".bmp": InputFormat.IMAGE,
    ".tiff": InputFormat.IMAGE,
    ".tif": InputFormat.IMAGE,
}


def _resolve_accelerator_device(device: str) -> AcceleratorDevice:
    """Resolve a user-provided device string to a Docling AcceleratorDevice."""
    normalized = (device or "auto").strip().lower()
    device_map = {
        "auto": AcceleratorDevice.AUTO,
        "cpu": AcceleratorDevice.CPU,
        "cuda": getattr(AcceleratorDevice, "CUDA", AcceleratorDevice.AUTO),
        "mps": getattr(AcceleratorDevice, "MPS", AcceleratorDevice.AUTO),
    }
    return device_map.get(normalized, AcceleratorDevice.AUTO)


def _build_table_structure_options(
    *,
    enable_table_merging: bool,
    enable_table_reconstruction: bool,
) -> TableStructureOptions:
    """Build table structure options, including optional flags when supported."""
    kwargs: dict[str, Any] = {"do_cell_matching": True}
    try:
        signature = inspect.signature(TableStructureOptions)
    except (TypeError, ValueError):
        return TableStructureOptions(**kwargs)

    if "do_table_merging" in signature.parameters:
        kwargs["do_table_merging"] = enable_table_merging
    if "do_table_reconstruction" in signature.parameters:
        kwargs["do_table_reconstruction"] = enable_table_reconstruction

    return TableStructureOptions(**kwargs)


def _build_page_range(max_pages_value: int | None) -> Any | None:
    """Build a Docling PageRange for limiting pages, when supported."""
    if not max_pages_value:
        return None
    if PageRange is None:
        return (1, max_pages_value)
    try:
        signature = inspect.signature(PageRange)
        if "start" in signature.parameters or "end" in signature.parameters:
            return PageRange(start=1, end=max_pages_value)
        return PageRange(1, max_pages_value)
    except Exception:
        try:
            return PageRange(1, max_pages_value)
        except Exception:
            return (1, max_pages_value)


@dataclass
class ExtractedImage:
    """An image extracted from a document."""

    index: int  # 1-based index
    pil_image: PILImage.Image
    caption: str | None = None
    page_number: int | None = None


@st.cache_resource(show_spinner=False)
def _get_docling_converter(
    enable_ocr: bool,
    enable_table_structure: bool,
    num_threads: int,
    device: str = "auto",
    enable_table_merging: bool = True,
    enable_table_reconstruction: bool = True,
    use_native_description: bool = False,
) -> DocumentConverter:
    """Create a cached docling converter tuned for speed."""
    safe_threads = max(1, int(num_threads or 4))

    pipeline_options = ThreadedPdfPipelineOptions(
        accelerator_options=AcceleratorOptions(
            num_threads=safe_threads,
            device=_resolve_accelerator_device(device),
        ),
        do_ocr=enable_ocr,
        do_table_structure=enable_table_structure,
        table_structure_options=_build_table_structure_options(
            enable_table_merging=enable_table_merging,
            enable_table_reconstruction=enable_table_reconstruction,
        ),
        # Disable non-essential outputs for speed
        generate_page_images=False,
        generate_table_images=False,
        generate_picture_images=False,
        generate_parsed_pages=False,
        do_picture_classification=use_native_description,
        do_picture_description=use_native_description,
    )

    format_option = FormatOption(
        pipeline_options=pipeline_options,
        backend=DoclingParseV4DocumentBackend,
        pipeline_cls=StandardPdfPipeline,
    )

    return DocumentConverter(
        allowed_formats=[InputFormat.PDF],
        format_options={InputFormat.PDF: format_option},
    )


@st.cache_resource(show_spinner=False)
def _get_generic_docling_converter(input_format: InputFormat) -> DocumentConverter:
    """Create a cached docling converter for non-PDF formats.

    Args:
        input_format: The InputFormat to convert (DOCX, PPTX, XLSX, HTML, IMAGE)

    Returns:
        Configured DocumentConverter for the specified format
    """
    return DocumentConverter(allowed_formats=[input_format])


@st.cache_resource(show_spinner=False)
def _get_docling_converter_with_images(
    enable_ocr: bool,
    enable_table_structure: bool,
    num_threads: int,
    device: str = "auto",
    enable_table_merging: bool = True,
    enable_table_reconstruction: bool = True,
    use_native_description: bool = False,
) -> DocumentConverter:
    """Create a cached docling converter with image extraction enabled."""
    safe_threads = max(1, int(num_threads or 4))

    pipeline_options = ThreadedPdfPipelineOptions(
        accelerator_options=AcceleratorOptions(
            num_threads=safe_threads,
            device=_resolve_accelerator_device(device),
        ),
        do_ocr=enable_ocr,
        do_table_structure=enable_table_structure,
        table_structure_options=_build_table_structure_options(
            enable_table_merging=enable_table_merging,
            enable_table_reconstruction=enable_table_reconstruction,
        ),
        # Enable image extraction
        generate_page_images=True,
        generate_table_images=False,
        generate_picture_images=True,
        generate_parsed_pages=False,
        do_picture_classification=use_native_description,
        do_picture_description=use_native_description,
    )

    format_option = FormatOption(
        pipeline_options=pipeline_options,
        backend=DoclingParseV4DocumentBackend,
        pipeline_cls=StandardPdfPipeline,
    )

    return DocumentConverter(
        allowed_formats=[InputFormat.PDF],
        format_options={InputFormat.PDF: format_option},
    )


def _extract_images_from_docling(doc: Any) -> list[ExtractedImage]:
    """Extract images from a docling document result."""
    images: list[ExtractedImage] = []
    index = 1

    for item, _level in doc.iterate_items():
        if isinstance(item, PictureItem):
            try:
                pil_image = item.get_image(doc)
                if pil_image is not None:
                    # Try to get page number from item's provenance
                    page_number = None
                    if hasattr(item, "prov") and item.prov:
                        for prov in item.prov:
                            if hasattr(prov, "page_no"):
                                page_number = prov.page_no
                                break

                    images.append(
                        ExtractedImage(
                            index=index,
                            pil_image=pil_image,
                            caption=None,
                            page_number=page_number,
                        )
                    )
                    index += 1
            except Exception:
                # Skip images that fail to extract
                continue

    return images


def export_document(
    doc: Any,
    output_format: str,
    filter_labels: list[DocItemLabel] | None = None,
) -> str:
    """Export a Docling document to the specified format.

    Args:
        doc: Docling document object
        output_format: Target format (markdown, html, doctags, json)
        filter_labels: Optional list of DocItemLabel types to filter out (markdown only)

    Returns:
        Exported document text in the specified format
    """
    normalized_format = (output_format or "markdown").strip().lower()
    if normalized_format == "markdown":
        # Use native Docling filtering by computing the set of allowed labels
        if filter_labels:
            # Get all possible labels and exclude the filtered ones
            all_labels = set(DocItemLabel)
            allowed_labels = all_labels - set(filter_labels)
            return cast(str, doc.export_to_markdown(labels=allowed_labels))
        serialized = _serialize_docling_document(doc, normalized_format)
        if serialized is not None:
            return serialized
        return cast(str, doc.export_to_markdown())
    elif normalized_format == "html":
        serialized = _serialize_docling_document(doc, normalized_format)
        if serialized is not None:
            return serialized
        return cast(str, doc.export_to_html())
    elif normalized_format == "doctags":
        return cast(str, doc.export_to_document_tokens())
    elif normalized_format == "json":
        return cast(str, doc.model_dump_json())
    else:
        # Default to markdown
        if filter_labels:
            all_labels = set(DocItemLabel)
            allowed_labels = all_labels - set(filter_labels)
            return cast(str, doc.export_to_markdown(labels=allowed_labels))
        return cast(str, doc.export_to_markdown())


def _serialize_docling_document(doc: Any, output_format: str) -> str | None:
    """Serialize a Docling document using the document serializer API."""
    normalized = (output_format or "markdown").strip().lower()
    if normalized == "markdown" and MarkdownDocSerializer is not None:
        try:
            serializer = MarkdownDocSerializer(doc=doc)
            return serializer.serialize().text
        except Exception:
            return None
    if normalized == "html" and HTMLDocSerializer is not None:
        try:
            serializer = HTMLDocSerializer(doc=doc)
            return serializer.serialize().text
        except Exception:
            return None
    return None


def parse_pdf_docling(
    content: bytes, params: dict[str, Any] | None = None
) -> tuple[str, list[ExtractedImage]]:
    """Parse PDF content using Docling engine.

    Args:
        content: PDF file content as bytes
        params: Optional parsing configuration including:
            - output_format: Export format (markdown, html, doctags, json)
            - docling_enable_ocr: Enable OCR for scanned PDFs
            - docling_table_structure: Extract table structure
            - docling_threads: Number of worker threads
            - docling_device: Compute device (auto, cpu, cuda, mps)
            - docling_filter_labels: Labels to filter out (markdown only)
            - docling_extract_images: Extract images from PDF
            - max_pages: Maximum number of pages to parse

    Returns:
        Tuple of (exported_text, extracted_images)

    Raises:
        ValueError: If PDF parsing fails
    """
    params = params or {}

    # For large digital PDFs we skip OCR by default for speed, but allow opt-in.
    enable_ocr = bool(params.get("docling_enable_ocr", False))
    enable_table_structure = bool(params.get("docling_table_structure", True))
    num_threads = params.get("docling_threads") or os.cpu_count() or 4
    extract_images = bool(params.get("docling_extract_images", False))
    use_native_description = bool(params.get("docling_use_native_description", False))
    device = params.get("docling_device", "auto")

    # Advanced table options
    enable_table_merging = bool(params.get("enable_table_merging", True))
    enable_table_reconstruction = bool(params.get("enable_table_reconstruction", True))

    # Get output format (markdown, html, doctags, json)
    output_format = params.get("output_format", "markdown")

    # Build list of labels to filter from the list of label names (markdown only)
    filter_label_names = cast(list[str], params.get("docling_filter_labels", []))
    filter_labels: list[DocItemLabel] = []
    for label_name in filter_label_names:
        try:
            filter_labels.append(DocItemLabel(label_name.lower()))
        except ValueError:
            # Skip invalid label names
            pass

    max_pages = params.get("max_pages")
    max_pages_value = int(max_pages) if isinstance(max_pages, int) and max_pages > 0 else None
    page_range = _build_page_range(max_pages_value)
    page_range = _build_page_range(max_pages_value)

    try:
        import tempfile

        # docling requires a file path, so we'll use a temp file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:

            def _convert_pdf_with_docling(
                *,
                enable_ocr_value: bool,
                enable_table_structure_value: bool,
                num_threads_value: int,
                extract_images_value: bool,
                enable_table_merging_value: bool,
                enable_table_reconstruction_value: bool,
                use_native_description_value: bool,
            ) -> Any:
                if extract_images_value:
                    converter = _get_docling_converter_with_images(
                        enable_ocr=enable_ocr_value,
                        enable_table_structure=enable_table_structure_value,
                        num_threads=num_threads_value,
                        device=device,
                        enable_table_merging=enable_table_merging_value,
                        enable_table_reconstruction=enable_table_reconstruction_value,
                        use_native_description=use_native_description_value,
                    )
                else:
                    converter = _get_docling_converter(
                        enable_ocr=enable_ocr_value,
                        enable_table_structure=enable_table_structure_value,
                        num_threads=num_threads_value,
                        device=device,
                        enable_table_merging=enable_table_merging_value,
                        enable_table_reconstruction=enable_table_reconstruction_value,
                        use_native_description=use_native_description_value,
                    )
                if page_range:
                    result = converter.convert(tmp_path, page_range=page_range)
                else:
                    result = converter.convert(tmp_path)
                return result.document

            try:
                doc = _convert_pdf_with_docling(
                    enable_ocr_value=enable_ocr,
                    enable_table_structure_value=enable_table_structure,
                    num_threads_value=num_threads,
                    extract_images_value=extract_images,
                    enable_table_merging_value=enable_table_merging,
                    enable_table_reconstruction_value=enable_table_reconstruction,
                    use_native_description_value=use_native_description,
                )
            except Exception as primary_error:
                # Retry with safe settings for more robust parsing.
                try:
                    doc = _convert_pdf_with_docling(
                        enable_ocr_value=False,
                        enable_table_structure_value=False,
                        num_threads_value=1,
                        extract_images_value=extract_images,
                        enable_table_merging_value=False,
                        enable_table_reconstruction_value=False,
                        use_native_description_value=False,
                    )
                except Exception as safe_error:
                    raise ValueError(
                        "Failed to parse PDF with docling: "
                        f"{str(primary_error)}. "
                        "Retry with safe settings failed: "
                        f"{str(safe_error)}."
                    ) from safe_error

            # Export document in the specified format
            text = export_document(doc, output_format, filter_labels)

            # Extract images if enabled
            images: list[ExtractedImage] = []
            if extract_images:
                images = _extract_images_from_docling(doc)

            return text, images
        finally:
            # Clean up temp file
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass
    except Exception as e:
        raise ValueError(f"Failed to parse PDF with docling: {str(e)}") from e


def parse_pdf(
    content: bytes, params: dict[str, Any] | None = None
) -> tuple[str, list[ExtractedImage]]:
    """Parse PDF content using Docling.

    Args:
        content: PDF file content as bytes
        params: Optional parsing configuration

    Returns:
        Tuple of (text_content, extracted_images)

    Raises:
        ValueError: If PDF parsing fails
    """
    return parse_pdf_docling(content, params)


def parse_with_docling(
    content: bytes,
    extension: str,
    params: dict[str, Any] | None = None,
) -> tuple[str, list[ExtractedImage]]:
    """
    Parse a non-PDF document using Docling and return exported text and any extracted images.
    
    Parameters:
        content (bytes): File content.
        extension (str): File extension including leading dot (e.g., ".pptx", ".xlsx", ".html", ".png").
        params (dict[str, Any] | None): Optional settings:
            - output_format: "markdown", "html", "doctags", or "json" (default "markdown").
            - docling_filter_labels: list[str] of label names to include when exporting markdown.
            - docling_extract_images: bool to enable image extraction.
            - max_pages: positive int to limit pages parsed.
    
    Returns:
        tuple[str, list[ExtractedImage]]: Exported document text and a list of extracted images.
    
    Raises:
        ValueError: If the file extension is not supported by Docling or if parsing fails.
    """
    params = params or {}
    extension = extension.lower()

    if extension not in DOCLING_FORMAT_MAP:
        raise ValueError(f"Unsupported format for Docling: {extension}")

    input_format = DOCLING_FORMAT_MAP[extension]

    # Get output format (markdown, html, doctags, json)
    output_format = params.get("output_format", "markdown")

    # Build list of labels to filter from the list of label names (markdown only)
    filter_label_names = cast(list[str], params.get("docling_filter_labels", []))
    filter_labels: list[DocItemLabel] = []
    for label_name in filter_label_names:
        try:
            filter_labels.append(DocItemLabel(label_name.lower()))
        except ValueError:
            pass

    max_pages = params.get("max_pages")
    max_pages_value = int(max_pages) if isinstance(max_pages, int) and max_pages > 0 else None
    page_range = _build_page_range(max_pages_value)

    try:
        import tempfile

        # Docling requires a file path
        with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            converter = _get_generic_docling_converter(input_format)
            if page_range:
                result = converter.convert(tmp_path, page_range=page_range)
            else:
                result = converter.convert(tmp_path)
            doc = result.document

            # Export document in the specified format
            text = export_document(doc, output_format, filter_labels)

            # Extract images if present
            images: list[ExtractedImage] = []
            if params.get("docling_extract_images", False):
                images = _extract_images_from_docling(doc)

            return text, images
        finally:
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass
    except Exception as e:
        format_name = extension.lstrip(".").upper()
        raise ValueError(f"Failed to parse {format_name} with Docling: {str(e)}") from e


def _extract_table_as_markdown(table: Any) -> str:
    """Extract DOCX table as markdown-style text.

    Args:
        table: python-docx Table object

    Returns:
        Markdown-formatted table text
    """
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))

    # Add separator after header row
    if rows:
        header = rows[0]
        separator = " | ".join(["---"] * len(rows[0].split(" | ")))
        return f"{header}\n{separator}\n" + "\n".join(rows[1:])
    return ""


def parse_docx(content: bytes, params: dict[str, Any] | None = None) -> str:
    """Parse DOCX content with optional markdown conversion.

    Args:
        content: DOCX file content as bytes
        params: Optional parsing configuration

    Returns:
        Extracted text content (optionally as markdown)

    Raises:
        ImportError: If python-docx is not installed
        ValueError: If DOCX parsing fails
    """
    params = params or {}
    if Document is None:
        raise ImportError(
            "python-docx is required for DOCX parsing. " "Install with: pip install python-docx"
        )

    try:
        from io import BytesIO

        docx_file = BytesIO(content)
        doc = Document(docx_file)
        text_parts = []

        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Process tables if enabled
        if params.get("extract_tables", True) and doc.tables:
            for table in doc.tables:
                table_text = _extract_table_as_markdown(table)
                if table_text:
                    text_parts.append(f"\n\n{table_text}\n")

        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}") from e


def parse_markdown(content: bytes, params: dict[str, Any] | None = None) -> str:
    """Parse Markdown content.

    Args:
        content: Markdown file content as bytes
        params: Optional parsing configuration

    Returns:
        Extracted text content (raw markdown or plain text based on params)

    Raises:
        ImportError: If markdown is not installed
        ValueError: If Markdown parsing fails
    """
    params = params or {}

    try:
        text = content.decode("utf-8")

        # If output format is markdown or original, return raw markdown
        output_format = params.get("output_format", "original")
        if output_format in ["markdown", "original"]:
            return text

        # Otherwise convert to plain text
        if markdown is None:
            # If markdown library not available, just return text as-is
            return text

        # Convert markdown to HTML, then strip HTML tags for plain text
        html = markdown.markdown(text)
        # Simple HTML tag removal (basic implementation)
        import re

        plain_text = re.sub(r"<[^>]+>", "", html)
        # Decode HTML entities
        import html as html_lib

        return html_lib.unescape(plain_text)
    except Exception as e:
        raise ValueError(f"Failed to parse Markdown: {str(e)}") from e


def parse_text(content: bytes) -> str:
    """Parse plain text content.

    Args:
        content: Text file content as bytes

    Returns:
        Text content as string
    """
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        # Try other common encodings
        for encoding in ["latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Failed to decode text file with common encodings") from None


def _markdown_to_plain_text(text: str) -> str:
    """Convert markdown content to plain text."""
    import html as html_lib
    import re

    # If markdown lib available, strip tags via HTML conversion for better fidelity
    if markdown is not None:
        html_text = markdown.markdown(text)
        plain_text = re.sub(r"<[^>]+>", "", html_text)
        plain_text = re.sub(r"\s+\n", "\n", plain_text)
        return html_lib.unescape(plain_text)

    # Fallback: lightweight regex cleanup
    text = re.sub(r"^#+\s", "", text, flags=re.MULTILINE)  # Remove headings
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # Remove bold
    text = re.sub(r"_([^_]+)_", r"\1", text)  # Remove italic
    text = re.sub(r"\|[^\n]+\|", "", text)  # Remove table rows
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)  # Remove separators
    text = re.sub(r"^\s*[-*]\s+", "", text, flags=re.MULTILINE)  # Remove list bullets
    return html_lib.unescape(text)


def _convert_to_format(text: str, source_format: str, target_format: str) -> str:  # noqa: ARG001
    """Convert text between different formats.

    For Docling-parsed documents (PDF, PPTX, XLSX, HTML, Image), the output
    format is already applied during parsing via export_document().

    For non-Docling parsers (DOCX, Markdown, Text), this function handles
    format conversion where possible.

    Args:
        text: Input text
        source_format: Source format (PDF, DOCX, Markdown, Text, etc.)
        target_format: Target format (markdown, html, doctags, json)

    Returns:
        Converted text
    """
    # New Docling formats - text is already in the correct format from parsing
    if target_format in ("markdown", "html", "doctags", "json"):
        return text
    # Legacy formats for backwards compatibility
    elif target_format == "original":
        return text
    elif target_format == "plain_text":
        return _markdown_to_plain_text(text)
    else:
        return text


def parse_document(
    filename: str,
    content: bytes,
    params: dict[str, Any] | None = None,
) -> tuple[str, str, list[ExtractedImage]]:
    """Parse a document based on its file extension.

    Args:
        filename: Name of the file (used to determine format)
        content: File content as bytes
        params: Optional parsing configuration dictionary

    Returns:
        Tuple of (parsed_text, file_format, extracted_images)

    Raises:
        ValueError: If file format is not supported or parsing fails
    """
    params = params or {}
    file_path = Path(filename)
    extension = file_path.suffix.lower()

    # Parse document based on format
    images: list[ExtractedImage] = []

    if extension == ".pdf":
        parsed_text, images = parse_pdf(content, params)
        file_format = "PDF"
    elif extension == ".docx":
        parsed_text = parse_docx(content, params)
        file_format = "DOCX"
    elif extension == ".pptx":
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "PPTX"
    elif extension == ".xlsx":
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "XLSX"
    elif extension in [".html", ".htm"]:
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "HTML"
    elif extension in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]:
        parsed_text, images = parse_with_docling(content, extension, params)
        file_format = "Image"
    elif extension in [".md", ".markdown"]:
        parsed_text = parse_markdown(content, params)
        file_format = "Markdown"
    elif extension == ".txt":
        parsed_text = parse_text(content)
        file_format = "Text"
    else:
        raise ValueError(
            f"Unsupported file format: {extension}. "
            f"Supported formats: .pdf, .docx, .pptx, .xlsx, .html, .md, .txt, .png, .jpg"
        )

    # Apply output format conversion (for non-Docling parsers)
    output_format = params.get("output_format", "markdown")
    parsed_text = _convert_to_format(parsed_text, file_format, output_format)

    return parsed_text, file_format, images


def get_file_preview(text: str, max_length: int = 500) -> str:
    """Get a preview of document text.

    Args:
        text: Full document text
        max_length: Maximum length of preview in characters

    Returns:
        Preview text (truncated if necessary)
    """
    if len(text) <= max_length:
        return text

    # Try to truncate at a sentence boundary
    preview = text[:max_length]
    last_period = preview.rfind(".")
    last_newline = preview.rfind("\n")

    # Use the later of period or newline, or just truncate
    truncate_at = max(last_period, last_newline)
    if truncate_at > max_length * 0.7:  # Only use boundary if it's not too early
        preview = preview[: truncate_at + 1]
    else:
        preview = preview[:max_length] + "..."

    return preview + "..." if len(text) > max_length else preview